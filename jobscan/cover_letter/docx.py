"""DOCX export for cover letters."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


class DocxExporter:
    """Export cover letters to DOCX format."""

    def __init__(
        self,
        font_name: str = "Georgia",
        font_size: int = 11,
    ) -> None:
        """Initialize the exporter.

        Args:
            font_name: Font to use for the document.
            font_size: Font size in points.
        """
        self.font_name = font_name
        self.font_size = font_size

    def export(
        self,
        content: str,
        output_dir: Path,
        company: str,
        position: str,
    ) -> Path:
        """Export cover letter to DOCX.

        Args:
            content: The cover letter text (plain text or markdown).
            output_dir: Directory to save the DOCX.
            company: Company name for the filename.
            position: Position title for the filename.

        Returns:
            Path to the generated DOCX file.
        """
        doc = Document()

        # Set default font for the document
        style = doc.styles['Normal']
        style.font.name = self.font_name
        style.font.size = Pt(self.font_size)

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Parse and add content
        self._add_content(doc, content)

        # Generate filename
        filename = self._generate_filename(company, position)
        output_path = output_dir / filename

        # Save document
        doc.save(output_path)

        return output_path

    def _add_content(self, doc: Document, content: str) -> None:
        """Add content to the document, handling basic markdown."""
        # Split into paragraphs
        paragraphs = content.strip().split('\n\n')

        for para_text in paragraphs:
            # Skip empty paragraphs
            if not para_text.strip():
                continue

            # Handle line breaks within paragraphs
            para_text = para_text.replace('\n', ' ').strip()

            # Create paragraph
            para = doc.add_paragraph()

            # Handle basic markdown bold/italic
            self._add_formatted_text(para, para_text)

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with basic markdown formatting (bold, italic)."""
        # Pattern to match **bold** and *italic*
        pattern = r'(\*\*.*?\*\*|\*.*?\*)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                # Regular text
                paragraph.add_run(part)

    def _generate_filename(self, company: str, position: str) -> str:
        """Generate a safe filename for the DOCX."""
        company_clean = self._sanitize_filename(company)
        position_clean = self._sanitize_filename(position)

        return f"cover_letter_{company_clean}_{position_clean}.docx"

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize a string for use in a filename."""
        # Replace spaces with underscores
        name = name.replace(" ", "_")
        # Remove any characters that aren't alphanumeric or underscores
        name = re.sub(r"[^\w]", "", name)
        # Truncate to reasonable length
        return name[:50].lower()


class DocxExportError(Exception):
    """Error exporting to DOCX."""

    pass
